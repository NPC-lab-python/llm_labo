"""Service d'export des projets vers Word/Markdown."""

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger
from sqlalchemy.orm import Session

from app.models.database import Project, ProjectSection, ProjectSource, Document, Reference
from config import settings


class CitationFormatter:
    """Formateur de citations selon différents styles."""

    @staticmethod
    def format_apa(
        authors: str | None,
        year: int | None,
        title: str | None,
        journal: str | None = None,
        volume: str | None = None,
        pages: str | None = None,
        doi: str | None = None,
    ) -> str:
        """
        Formate une référence au style APA 7e édition.

        Format: Auteur, A. A., & Auteur, B. B. (Année). Titre de l'article. Nom du Journal, volume, pages. https://doi.org/xxx
        """
        parts = []

        # Auteurs
        if authors:
            # Essayer de parser le JSON, sinon utiliser tel quel
            try:
                author_list = json.loads(authors)
                if isinstance(author_list, list):
                    formatted_authors = []
                    for author in author_list[:7]:  # APA limite à 7 auteurs
                        if isinstance(author, dict):
                            name = author.get("name", "")
                        else:
                            name = str(author)
                        # Convertir "Prénom Nom" en "Nom, P."
                        name_parts = name.split()
                        if len(name_parts) >= 2:
                            surname = name_parts[-1]
                            initials = ". ".join([p[0] + "." for p in name_parts[:-1]])
                            formatted_authors.append(f"{surname}, {initials}")
                        else:
                            formatted_authors.append(name)

                    if len(author_list) > 7:
                        authors_str = ", ".join(formatted_authors[:6]) + ", ... " + formatted_authors[-1]
                    elif len(formatted_authors) == 2:
                        authors_str = " & ".join(formatted_authors)
                    elif len(formatted_authors) > 2:
                        authors_str = ", ".join(formatted_authors[:-1]) + ", & " + formatted_authors[-1]
                    else:
                        authors_str = formatted_authors[0] if formatted_authors else "Auteur inconnu"
                else:
                    authors_str = str(author_list)
            except (json.JSONDecodeError, TypeError):
                authors_str = authors
        else:
            authors_str = "Auteur inconnu"

        parts.append(authors_str)

        # Année
        if year:
            parts.append(f"({year})")
        else:
            parts.append("(s.d.)")  # sans date

        # Titre
        if title:
            parts.append(f"{title}.")

        # Journal, volume, pages
        if journal:
            journal_part = f"*{journal}*"
            if volume:
                journal_part += f", *{volume}*"
            if pages:
                journal_part += f", {pages}"
            journal_part += "."
            parts.append(journal_part)

        # DOI
        if doi:
            if not doi.startswith("http"):
                doi = f"https://doi.org/{doi}"
            parts.append(doi)

        return " ".join(parts)

    @staticmethod
    def format_apa_in_text(authors: str | None, year: int | None) -> str:
        """
        Formate une citation dans le texte au style APA.

        Format: (Auteur, Année) ou (Auteur & Auteur, Année)
        """
        author_str = "Auteur inconnu"
        if authors:
            try:
                author_list = json.loads(authors)
                if isinstance(author_list, list) and author_list:
                    if isinstance(author_list[0], dict):
                        names = [a.get("name", "").split()[-1] for a in author_list]
                    else:
                        names = [str(a).split()[-1] for a in author_list]

                    if len(names) == 1:
                        author_str = names[0]
                    elif len(names) == 2:
                        author_str = f"{names[0]} & {names[1]}"
                    else:
                        author_str = f"{names[0]} et al."
                else:
                    author_str = str(author_list).split()[-1]
            except (json.JSONDecodeError, TypeError):
                author_str = authors.split(",")[0].split()[-1] if authors else "Auteur inconnu"

        year_str = str(year) if year else "s.d."
        return f"({author_str}, {year_str})"


class ExportService:
    """Service d'export des projets."""

    def __init__(self):
        """Initialise le service d'export."""
        self.export_dir = settings.data_dir / "exports"
        self.export_dir.mkdir(parents=True, exist_ok=True)
        self.citation_formatter = CitationFormatter()

    def export_project(
        self,
        db: Session,
        project: Project,
        format: str = "docx",
        include_bibliography: bool = True,
        citation_style: str = "apa",
    ) -> Path:
        """
        Exporte un projet au format demandé.

        Args:
            db: Session SQLAlchemy.
            project: Projet à exporter.
            format: Format d'export (docx, markdown).
            include_bibliography: Inclure la bibliographie.
            citation_style: Style de citation (apa).

        Returns:
            Chemin vers le fichier exporté.
        """
        logger.info(f"Export du projet '{project.title}' au format {format}")

        if format == "docx":
            return self._export_docx(db, project, include_bibliography, citation_style)
        elif format == "markdown":
            return self._export_markdown(db, project, include_bibliography, citation_style)
        else:
            raise ValueError(f"Format non supporté: {format}")

    def _export_docx(
        self,
        db: Session,
        project: Project,
        include_bibliography: bool,
        citation_style: str,
    ) -> Path:
        """Exporte en format Word (.docx)."""
        doc = DocxDocument()

        # Style du titre
        title = doc.add_heading(project.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Description si présente
        if project.description:
            desc = doc.add_paragraph(project.description)
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Espace

        # Sections triées par ordre
        sections = sorted(project.sections, key=lambda s: s.section_order)

        section_titles = {
            "introduction": "Introduction",
            "literature_review": "Revue de littérature",
            "methods": "Méthodologie",
            "results": "Résultats",
            "discussion": "Discussion",
            "conclusion": "Conclusion",
        }

        for section in sections:
            # Titre de section
            section_title = section.title or section_titles.get(section.section_type, section.section_type.title())
            doc.add_heading(section_title, 1)

            # Contenu
            if section.content:
                # Traiter le markdown simple
                paragraphs = section.content.split("\n\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        # Convertir les références [Source N] en citations APA
                        processed_text = self._process_citations(para_text, project.sources, citation_style)
                        para = doc.add_paragraph(processed_text)

        # Bibliographie
        if include_bibliography and project.sources:
            doc.add_page_break()
            doc.add_heading("Références", 1)

            # Collecter et trier les références
            references = []
            for source in project.sources:
                if source.document:
                    ref_text = self.citation_formatter.format_apa(
                        authors=source.document.authors,
                        year=source.document.publication_year,
                        title=source.document.title,
                        journal=source.document.journal,
                        doi=source.document.doi,
                    )
                    references.append(ref_text)

            # Trier alphabétiquement
            references.sort()

            for ref in references:
                para = doc.add_paragraph(ref, style="List Bullet")

        # Sauvegarder
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', project.title)[:50]
        filename = f"{safe_title}_{timestamp}.docx"
        file_path = self.export_dir / filename

        doc.save(str(file_path))
        logger.info(f"Projet exporté vers: {file_path}")

        return file_path

    def _export_markdown(
        self,
        db: Session,
        project: Project,
        include_bibliography: bool,
        citation_style: str,
    ) -> Path:
        """Exporte en format Markdown."""
        lines = []

        # Titre
        lines.append(f"# {project.title}")
        lines.append("")

        # Description
        if project.description:
            lines.append(f"*{project.description}*")
            lines.append("")

        # Sections
        sections = sorted(project.sections, key=lambda s: s.section_order)

        section_titles = {
            "introduction": "Introduction",
            "literature_review": "Revue de littérature",
            "methods": "Méthodologie",
            "results": "Résultats",
            "discussion": "Discussion",
            "conclusion": "Conclusion",
        }

        for section in sections:
            section_title = section.title or section_titles.get(section.section_type, section.section_type.title())
            lines.append(f"## {section_title}")
            lines.append("")

            if section.content:
                processed_text = self._process_citations(section.content, project.sources, citation_style)
                lines.append(processed_text)
                lines.append("")

        # Bibliographie
        if include_bibliography and project.sources:
            lines.append("---")
            lines.append("")
            lines.append("## Références")
            lines.append("")

            references = []
            for source in project.sources:
                if source.document:
                    ref_text = self.citation_formatter.format_apa(
                        authors=source.document.authors,
                        year=source.document.publication_year,
                        title=source.document.title,
                        journal=source.document.journal,
                        doi=source.document.doi,
                    )
                    references.append(ref_text)

            references.sort()
            for ref in references:
                lines.append(f"- {ref}")

        # Sauvegarder
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', project.title)[:50]
        filename = f"{safe_title}_{timestamp}.md"
        file_path = self.export_dir / filename

        file_path.write_text("\n".join(lines), encoding="utf-8")
        logger.info(f"Projet exporté vers: {file_path}")

        return file_path

    def _process_citations(
        self,
        text: str,
        sources: list,
        citation_style: str,
    ) -> str:
        """
        Remplace les marqueurs [Source N] par des citations formatées.

        Args:
            text: Texte avec marqueurs.
            sources: Liste des sources du projet.
            citation_style: Style de citation.

        Returns:
            Texte avec citations formatées.
        """
        # Créer un mapping source_index -> citation
        source_map = {}
        for i, source in enumerate(sources, 1):
            if source.document:
                citation = self.citation_formatter.format_apa_in_text(
                    authors=source.document.authors,
                    year=source.document.publication_year,
                )
                source_map[i] = citation

        # Remplacer [Source N] par la citation APA
        def replace_citation(match):
            source_num = int(match.group(1))
            return source_map.get(source_num, match.group(0))

        processed = re.sub(r'\[Source (\d+)\]', replace_citation, text)
        return processed


# Instance globale
export_service = ExportService()
